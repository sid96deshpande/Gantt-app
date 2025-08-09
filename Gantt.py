import openai
import os
import json
import re
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from datetime import datetime
import streamlit as st
import PyPDF2

# Optional: robust PDF table extraction
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# 🔐 API key
openai.api_key = st.secrets["openai_api_key"]

# 📀 Excel config
TASK_ROW_START = 8
TASK_ROW_GAP = 2
TASK_COLS = {'task': 4, 'assigned_to': 5, 'start': 6, 'end': 7}
START_DATE_CELL = 'F3'       # Project Start Date
PROJECT_NAME_CELL = 'F2'     # Project Name (E2 is label)

# -------------------- DOC INGEST --------------------
def extract_tables_from_docx(doc):
    """Extract all tables from a docx Document as CSV-style text."""
    tables_text = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(", ".join(cells))
        tables_text.append("\n".join(rows))
    return "\n\n".join(tables_text)

def read_scope_docx(docx_path):
    doc = Document(docx_path)
    paragraphs = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    tables = extract_tables_from_docx(doc)
    combined = paragraphs
    if tables:
        combined += "\n\n[TABLES]\n" + tables
    return combined

def extract_tables_from_pdf(pdf_path):
    """Try extracting tables from PDF using pdfplumber; fallback to empty."""
    if HAS_PDFPLUMBER:
        tables = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    rows = []
                    for row in table:
                        if row:
                            cells = [str(cell).strip().replace("\n", " ") if cell else "" for cell in row]
                            rows.append(", ".join(cells))
                    if rows:
                        tables.append("\n".join(rows))
        return "\n\n".join(tables)
    return ""

def read_scope_pdf(pdf_path):
    text = ""
    with open(pdf_path, "rb") as f:
        pdf = PyPDF2.PdfReader(f)
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    tables = extract_tables_from_pdf(pdf_path)
    combined = text
    if tables:
        combined += "\n\n[TABLES]\n" + tables
    return combined

# -------------------- HELPERS --------------------
def extract_json_from_response(content):
    match = re.search(r'\{[\s\S]*\}', content)
    if match:
        return match.group(0)
    return content

def parse_task_number(number):
    """Convert task number to numeric (int/float) if possible, else leave as str."""
    try:
        if isinstance(number, (int, float)):
            return number
        s = str(number)
        if "." in s:
            return float(s)
        return int(s)
    except Exception:
        return number

def parse_date(date_str):
    """Parse date to date() in multiple formats. Return None if invalid."""
    if not date_str or date_str == "null":
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d %b %Y", "%d %B %Y"):
        try:
            return datetime.strptime(date_str, fmt).date()
        except Exception:
            continue
    nums = re.findall(r"\d+", str(date_str))
    if len(nums) == 3:
        d, m, y = nums
        try:
            return datetime(int(y), int(m), int(d)).date()
        except Exception:
            pass
    return None

def sanitize_title(s: str) -> str:
    """Clean up a detected title to a short, single-line name."""
    if not s:
        return ""
    s = s.replace("\n", " ").strip()
    s = re.sub(r'\s+', ' ', s)
    return s[:120]

def infer_project_name(scope_text: str) -> str:
    """
    Fallback if LLM doesn't provide project_name:
    1) Line after 'Title'
    2) First non-empty line before 'Table of Contents'
    3) First meaningful non-empty line
    """
    lines = [ln.strip() for ln in scope_text.splitlines() if ln.strip()]
    for i, ln in enumerate(lines):
        if re.fullmatch(r'(?i)title', ln) and i + 1 < len(lines):
            cand = sanitize_title(lines[i + 1])
            if cand:
                return cand
    for i, ln in enumerate(lines):
        if re.search(r'(?i)table of contents', ln):
            for j in range(i - 1, -1, -1):
                cand = sanitize_title(lines[j])
                if cand:
                    return cand
            break
    if lines:
        return sanitize_title(lines[0])
    return "Project"

def to_number(val):
    """Parse numbers like '£1,200.50' or '1,200' -> 1200.5; return None if not numeric."""
    if isinstance(val, (int, float)):
        return float(val)
    if val is None:
        return None
    s = re.sub(r'[^\d.\-]', '', str(val))
    try:
        return float(s)
    except Exception:
        return None

# -------------------- LLM EXTRACTION --------------------
def extract_tasks_with_gpt(scope_text):
    """
    Ask the LLM to extract project_name, project_start_date, tasks, per-task budget/estimate, and total budget if present.
    """
    prompt = f"""
You are an AI project agent.
Below is text+tables extracted from a project scope (Word/PDF).

Return ONLY valid minified JSON with keys:
- "project_name": short title of the project (string; if unclear, infer best concise title)
- "project_start_date": YYYY-MM-DD (if missing/unclear, use earliest detectable date; else null)
- "project_total_budget": number or null (if a single total project budget is specified anywhere)
- "tasks": array of objects with:
  - "task": number/id (e.g., 1, 2, 1.1)
  - "description": short (<= 80 chars)
  - "assigned_to": string ("" if missing)
  - "start": YYYY-MM-DD or null
  - "end": YYYY-MM-DD or null
  - "estimated_cost": number or null (explicit estimate for this task, if given)
  - "budget": number or null (explicit budget for this task, if given)

Notes:
- Budgets/estimates can appear anywhere, including at the end; parse carefully from text or tables.
- Ranges like "12–18 June" => inclusive in year 2025 unless a year is explicit.
- If no dates for a task, set start/end to null.
- Emit ONLY JSON, no commentary.

Example:
{{"project_name":"Example Project","project_start_date":"2025-01-01","project_total_budget":120000,"tasks":[{{"task":1,"description":"Site survey","assigned_to":"Site Eng","start":"2025-01-01","end":"2025-01-02","estimated_cost":1200,"budget":1000}}]}}

Project Scope:
\"\"\"{scope_text}\"\"\"
"""
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0
    )
    content = response.choices[0].message.content.strip()
    json_str = extract_json_from_response(content)
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        st.error("❌ GPT returned invalid JSON. See console for details.")
        print("⚠️ GPT raw response:", content)
        raise

# -------------------- EXCEL WRITERS --------------------
def copy_row_style(source_row, target_row, ws):
    for col in range(2, 9):  # Columns B to H
        source = ws.cell(row=source_row, column=col)
        target = ws.cell(row=target_row, column=col)
        if source.has_style:
            target.font = source.font.copy()
            target.border = source.border.copy()
            target.fill = source.fill.copy()
            target.number_format = source.number_format
            target.alignment = source.alignment.copy()

def get_market_price(item):
    # Baseline market assumptions (same idea as before)
    prices = {
        "labour": {"hour": 8, "per_hr": 25, "total": 200},
        "material": {"units": 5, "per_unit": 50, "total": 250},
        "travel": 10,
        "equipment": 30,
        "fixed": 0,
        "misc": 15
    }
    return prices.get(item, {})

def _baseline_labour_material(costs):
    """Compute baseline totals for labour and materials from provided costs or defaults."""
    # Labour
    lab_total = None
    if isinstance(costs.get("labour"), dict):
        l = costs["labour"]
        hours = to_number(l.get("hours"))
        per_hr = to_number(l.get("per_hr"))
        total = to_number(l.get("total"))
        if total is not None:
            lab_total = total
        else:
            h = hours if hours is not None else get_market_price("labour").get("hour", 8)
            r = per_hr if per_hr is not None else get_market_price("labour").get("per_hr", 25)
            lab_total = h * r
    else:
        maybe_total = to_number(costs.get("labour"))
        if maybe_total is not None:
            lab_total = maybe_total
        else:
            lab_total = get_market_price("labour").get("hour", 8) * get_market_price("labour").get("per_hr", 25)

    # Materials
    mat_total = None
    if isinstance(costs.get("material"), dict):
        m = costs["material"]
        units = to_number(m.get("units"))
        per_unit = to_number(m.get("per_unit"))
        total = to_number(m.get("total"))
        if total is not None:
            mat_total = total
        else:
            u = units if units is not None else get_market_price("material").get("units", 5)
            pu = per_unit if per_unit is not None else get_market_price("material").get("per_unit", 50)
            mat_total = u * pu
    else:
        maybe_total = to_number(costs.get("material"))
        if maybe_total is not None:
            mat_total = maybe_total
        else:
            mat_total = get_market_price("material").get("units", 5) * get_market_price("material").get("per_unit", 50)

    return float(labour_total := lab_total), float(material_total := mat_total)

def _other_components(costs):
    """Return a dict of other components with numbers filled from scope or market defaults."""
    out = {}
    out["travel"] = to_number(costs.get("travel"))
    out["equipment"] = to_number(costs.get("equipment"))
    out["fixed"] = to_number(costs.get("fixed"))
    out["misc"] = to_number(costs.get("misc"))
    if out["travel"] is None:    out["travel"] = get_market_price("travel")
    if out["equipment"] is None: out["equipment"] = get_market_price("equipment")
    if out["fixed"] is None:     out["fixed"] = get_market_price("fixed")
    if out["misc"] is None:      out["misc"] = get_market_price("misc")
    return out

def _backsolve_pair(total, first=None, second=None, default_first=None, default_second=None, max_dp=4):
    """Given total and (first, second), compute a pair so first*second == total.
    If one of first/second is provided, compute the other. Otherwise use defaults.
    Values are rounded to reasonable decimals to ensure the product matches total closely.
    """
    t = float(total or 0)
    if t <= 0:
        return 0.0, 0.0

    # If both provided, adjust second to match exactly
    if first not in (None, 0) and second not in (None, 0):
        sec = round(t / float(first), max_dp)
        return float(first), sec

    # If only first provided
    if first not in (None, 0):
        sec = round(t / float(first), max_dp)
        return float(first), sec

    # If only second provided
    if second not in (None, 0):
        fir = round(t / float(second), max_dp)
        return fir, float(second)

    # Neither provided: use defaults then backsolve
    f = default_first if default_first not in (None, 0) else 1.0
    s = round(t / float(f), max_dp)
    return float(f), s

def fill_gantt_excel(template_path, output_path, project_name, start_date, tasks):
    wb = load_workbook(template_path)
    ws = wb.active

    # Project Name -> F2
    ws[PROJECT_NAME_CELL] = sanitize_title(project_name) or "Project"

    # Project Start Date -> F3
    sd = parse_date(start_date) or datetime.today().date()
    ws[START_DATE_CELL] = sd

    for i, task in enumerate(tasks):
        row = TASK_ROW_START + i * TASK_ROW_GAP
        if row > 26:
            copy_row_style(26, row, ws)

        ws.cell(row=row, column=1).value = parse_task_number(task.get("task", i + 1))
        ws.cell(row=row, column=2).value = 0
        ws.cell(row=row, column=TASK_COLS['task']).value = task.get("description", "")
        ws.cell(row=row, column=TASK_COLS['assigned_to']).value = task.get("assigned_to", "")

        start = parse_date(task.get("start")) or sd
        end = parse_date(task.get("end")) or start

        ws.cell(row=row, column=TASK_COLS['start']).value = start
        ws.cell(row=row, column=TASK_COLS['end']).value = end

        col_start = get_column_letter(TASK_COLS['start'])
        col_end = get_column_letter(TASK_COLS['end'])
        ws.cell(row=row, column=8).value = f"={col_end}{row}-{col_start}{row}"

    wb.save(output_path)

def fill_budget_excel(template_path, output_path, project_name, start_date, tasks, total_budget=None):
    wb = load_workbook(template_path)
    ws = wb.active

    # Keep existing headers usage
    sd = parse_date(start_date) or datetime.today().date()
    ws["C3"] = sanitize_title(project_name) or "Project"
    ws["D3"] = sd

    n_tasks = len(tasks)
    any_per_task_budget = False
    wrote_any_budget = False

    for i, task in enumerate(tasks):
        row = 8 + i
        costs = task.get("costs", {}) or {}
        description = task.get("description", "")
        ws[f"C{row}"] = description  # Description

        start = parse_date(task.get("start")) or sd
        ws[f"E{row}"] = start       # Planned Start
        end = parse_date(task.get("end")) or None
        ws[f"G{row}"] = end         # End Date

        # ---- BASELINES ----
        labour_base, material_base = _baseline_labour_material(costs)
        others = _other_components(costs)
        other_total = float(others["travel"]) + float(others["equipment"]) + float(others["fixed"]) + float(others["misc"])

        # Write other components (N..Q)
        ws[f"N{row}"] = others["travel"]
        ws[f"O{row}"] = others["equipment"]
        ws[f"P{row}"] = others["fixed"]
        ws[f"Q{row}"] = others["misc"]

        # Estimated task cost from scope (if any)
        estimated = to_number(task.get("estimated_cost"))
        if estimated is None:
            estimated = to_number(task.get("budget")) or to_number(task.get("est_cost")) or to_number(task.get("cost"))

        # ---- ALLOCATION ----
        if estimated is not None and estimated >= 0:
            # Keep other_total as-is; split remainder proportionally to baselines
            lm_base_sum = max(labour_base + material_base, 1e-6)
            remainder = max(0.0, float(estimated) - other_total)
            labour_final = round(remainder * (labour_base / lm_base_sum), 2)
            material_final = round(remainder - labour_final, 2)  # exact sum

            # Totals
            ws[f"J{row}"] = labour_final       # Labour Total
            ws[f"M{row}"] = material_final     # Materials Total
            ws[f"T{row}"] = float(estimated)   # Actual = estimated
        else:
            # No estimated: use baselines + others
            labour_final = round(labour_base, 2)
            material_final = round(material_base, 2)
            actual_total = round(labour_final + material_final + other_total, 2)

            ws[f"J{row}"] = labour_final
            ws[f"M{row}"] = material_final
            ws[f"T{row}"] = actual_total

        # ---- Back-solve Hours & £/Hr so H*I == J ----
        hours_in = None
        rate_in = None
        if isinstance(costs.get("labour"), dict):
            hours_in = to_number(costs["labour"].get("hours"))
            rate_in = to_number(costs["labour"].get("per_hr"))
        default_hours = get_market_price("labour").get("hour", 8)
        default_rate = get_market_price("labour").get("per_hr", 25)
        hours, rate = _backsolve_pair(labour_final, hours_in, rate_in, default_hours, default_rate, max_dp=4)
        ws[f"H{row}"] = hours
        ws[f"I{row}"] = rate

        # ---- Back-solve Units & £/Unit so K*L == M ----
        units_in = None
        unit_price_in = None
        if isinstance(costs.get("material"), dict):
            units_in = to_number(costs["material"].get("units"))
            unit_price_in = to_number(costs["material"].get("per_unit"))
        default_units = get_market_price("material").get("units", 5)
        default_unit_price = get_market_price("material").get("per_unit", 50)
        units, unit_price = _backsolve_pair(material_final, units_in, unit_price_in, default_units, default_unit_price, max_dp=4)
        ws[f"K{row}"] = units
        ws[f"L{row}"] = unit_price

        # ---- Budget column (S) per new rules ----
        per_task_budget = to_number(task.get("budget"))
        if per_task_budget is not None:
            any_per_task_budget = True
            ws[f"S{row}"] = per_task_budget
            wrote_any_budget = True
        elif total_budget is not None and n_tasks > 0:
            equal_share = round(float(total_budget) / n_tasks, 2)
            ws[f"S{row}"] = equal_share
            wrote_any_budget = True
        # else: leave blank

    wb.save(output_path)

    # Return whether any budget was written (either per-task or distributed total)
    return wrote_any_budget or any_per_task_budget or (total_budget is not None)

# -------------------- CLEANUP --------------------
def cleanup_temp_files():
    for temp_file in [
        "uploaded_scope.docx", "uploaded_scope.pdf",
        "filled_gantt.xlsx", "filled_budget.xlsx"
    ] + [fname for fname in os.listdir() if fname.startswith("filled_gantt_") or fname.startswith("filled_budget_")]:
        if os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except Exception:
                pass

# -------------------- ORCHESTRATOR --------------------
def run_agent(scope_text, gantt_template, budget_template):
    data = extract_tasks_with_gpt(scope_text)

    # Ensure we ALWAYS have a project name
    project_name = sanitize_title(data.get("project_name", "")) or infer_project_name(scope_text) or "Project"
    project_start_date = data.get("project_start_date")
    tasks = data.get("tasks", [])
    total_budget = to_number(data.get("project_total_budget"))

    gantt_out = f"filled_gantt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    budget_out = f"filled_budget_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    fill_gantt_excel(gantt_template, gantt_out, project_name, project_start_date, tasks)
    budgets_written = fill_budget_excel(budget_template, budget_out, project_name, project_start_date, tasks, total_budget=total_budget)

    return gantt_out, budget_out, budgets_written

# -------------------- STREAMLIT UI --------------------
st.set_page_config(page_title="Proja: Project Planning AI Agent", page_icon="📊", layout="centered")
st.image("Project_management.png", use_container_width=True)
st.title("Sid AI - Project Planning & Budget AI Agent")
st.info(
    "🔒 **Data Privacy:** Your uploaded files and outputs are deleted immediately after you download your results. "
    "No project data is stored or reused. "
    "Please avoid uploading files with sensitive personal information. "
    "[OpenAI Data Usage Policy](https://platform.openai.com/docs/data-usage-policies)"
)
st.markdown("---")
st.markdown(
    """
    <div style='background-color: #f3f6fa; padding: 16px; border-radius: 16px;'>
    <b>How to use:</b>
    <ul>
      <li>Upload your project scope <b>.docx</b>, <b>.pdf</b> <i>or</i> type your project details below.</li>
      <li>Get a Gantt Chart and Budget Excel – ready for download.</li>
      <li>Agent will only answer project breakdown, Gantt chart, or budget requests.</li>
    
      <li><b>This is merely a tool and is capable of producing errors. Users/project managers must use their discretion and review the documents carefully especially documents related to costs/budgeting.</b></li>
      </ul>
    </div>
    """,
    unsafe_allow_html=True,
)
st.sidebar.image("Sid.png", width=150)

with st.chat_message("assistant"):
    st.write("👋 Hi! I can help you create project Gantt charts and budgets. Please upload your project scope document (.docx, .pdf) or type your project details below.")

uploaded_file = st.file_uploader("Upload Project Scope (.docx, .pdf)", type=["docx", "pdf"])
user_text = st.chat_input("Or type your project scope / questions here...")

gantt_template = "Gantt Chart Template.xlsx"
budget_template = "Budget template.xlsx"

def handle_scope_input(scope_input):
    try:
        gantt_path, budget_path, budgets_written = run_agent(scope_input, gantt_template, budget_template)
        with st.container():
            st.markdown("#### 📦 Download Results")
            with open(gantt_path, "rb") as f1:
                st.download_button(
                    label="📅 Download Gantt Chart Excel",
                    data=f1,
                    file_name=gantt_path,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            with open(budget_path, "rb") as f2:
                st.download_button(
                    label="💷 Download Budget Excel",
                    data=f2,
                    file_name=budget_template,  # left unchanged per your request
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        if not budgets_written:
            st.warning("Please enter budget for each task")
        st.success("✅ Files created and ready for download.")
        cleanup_temp_files()
    except Exception as e:
        st.error(f"❌ An error occurred: {e}")
        cleanup_temp_files()

if uploaded_file is not None:
    if uploaded_file.name.endswith(".docx"):
        temp_docx_path = "uploaded_scope.docx"
        with open(temp_docx_path, "wb") as f:
            f.write(uploaded_file.read())
        with st.spinner("⏳ Analyzing your project scope and generating files..."):
            doc_text = read_scope_docx(temp_docx_path)
            handle_scope_input(doc_text)
    elif uploaded_file.name.endswith(".pdf"):
        temp_pdf_path = "uploaded_scope.pdf"
        with open(temp_pdf_path, "wb") as f:
            f.write(uploaded_file.read())
        with st.spinner("⏳ Analyzing your project scope and generating files..."):
            pdf_text = read_scope_pdf(temp_pdf_path)
            handle_scope_input(pdf_text)
elif user_text:
    keywords = ["gantt", "project", "budget", "tasks", "excel", "plan", "scope"]
    if not any(word in user_text.lower() for word in keywords):
        with st.chat_message("assistant"):
            st.write("Sorry, I can only help with generating project plans, Gantt charts, and budget sheets from your scope.")
    else:
        with st.spinner("⏳ Analyzing your project scope and generating files..."):
            handle_scope_input(user_text)
